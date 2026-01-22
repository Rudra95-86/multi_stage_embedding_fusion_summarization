import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import re
from io import BytesIO
from docx import Document
import PyPDF2

# Import logic from your modules
from preprocessing import preprocess_text
from word_embeddings import get_word2vec_embeddings
from sentence_embeddings import get_sentence_bert_embeddings
from contextual_embeddings import get_bert_embeddings
from textrank_summarizer import textrank_summary
from fusion import fuse_embeddings
from evaluation import similarity_score, get_rouge_scores

# --- PAGE CONFIG ---
st.set_page_config(page_title="Fusion AI Summarizer", layout="wide")

st.title("Multi-Stage Embedding Fusion Summarizer")
st.caption("Advanced Extraction using Unified Representation of Local, Contextual, and Thematic Data")
st.markdown("---")

# --- SIDEBAR SETTINGS ---
with st.sidebar:
    st.header("Model Parameters")
    summary_ratio = st.slider("Compression Ratio", 0.1, 0.8, 0.3)
    st.divider()
    st.markdown("System Architecture")
    with st.expander("How it works", expanded=True):
        st.write("""
        This uses **Multi-Stage Fusion**:
        - **Stage 1:** Word2Vec (captures local word-level semantics)
        - **Stage 2:** S-BERT (provides deep sentence-level context)
        - **Stage 3:** BERT (understands global thematic info)
        """)
    
    st.divider()
    
    # 2. Summary Tips
    st.markdown("Quick Tips")
    st.caption("- **Higher Ratio:** More detail, longer summary.")
    st.caption("- **Lower Ratio:** More concise, only core meaning.")

# --- INPUT SECTION ---
st.subheader("Source")
uploaded_file = st.file_uploader("Upload document", type=["txt", "docx","pdf"],width="stretch")

raw_text = ""
if uploaded_file:
    if uploaded_file.type == "text/plain":
        raw_text = str(uploaded_file.read(), "utf-8")
        
    elif:
        doc = Document(uploaded_file)
        raw_text = "\n".join([p.text for p in doc.paragraphs])\
        
    else:
        uploaded_file.type == "application/pdf"
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            content = page.extract_text()
            if content:
                raw_text += content + "\n"

# Manual input fallback
if not raw_text:
    raw_text = st.text_area("Or Paste Source Document Here...", height=200)

# --- VALIDATION LOGIC ---
is_valid = False
if raw_text:
    words = raw_text.split()
    word_count = len(words)
    
    # Check if input is purely numerical
    is_pure_numbers = all(re.match(r'^[0-9\W]+$', w) for w in words)
    
    if is_pure_numbers:
        st.error("❌ **Invalid Input:** The document contains only numbers/symbols. Please provide text for analysis.")
    elif word_count < 100:
        st.warning(f"⚠️ **Word Count Low:** Only {word_count} words detected. A minimum of 100 words is required.")
    else:
        st.info(f"✅ **Document Ready:** {word_count} words detected.")
        is_valid = True

# --- EXECUTION ---
if st.button("Generate Summary", width="stretch", disabled=not is_valid):
    with st.spinner('Fusing multi-stage embeddings...'):
        # 1. Pipeline
        sentences, tokenized_sentences = preprocess_text(raw_text)
        word_emb = get_word2vec_embeddings(tokenized_sentences)
        sent_emb = get_sentence_bert_embeddings(sentences)
        bert_emb = get_bert_embeddings(sentences)
        
        # 2. Fusion & Summarization
        fused_emb = fuse_embeddings(word_emb, sent_emb, bert_emb)
        summary_fused = textrank_summary(sentences, fused_emb, summary_ratio=summary_ratio)
        summary_single = textrank_summary(sentences, sent_emb, summary_ratio=summary_ratio)

        # --- TABS INTERFACE ---
        tab1, tab2 = st.tabs(["📄 Final Summary", "📊 Performance Analysis"])

        with tab1:
            st.subheader("Summary")
            # Native st.write ensures high contrast (Black text on Light mode, White on Dark)
            st.write(summary_fused)
            
            # DOCX Download
            doc_out = Document()
            doc_out.add_paragraph(summary_fused)
            bio = BytesIO()
            doc_out.save(bio)
            st.download_button("📥 Download Summary", bio.getvalue(), "summary.docx", 
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",width='stretch')

        # ... (Previous imports and initialization code remain the same)

        with tab2:
            st.subheader("Compression & Accuracy Evaluation")
            
            # 1. Performance Statistics Row
            summary_words = len(summary_fused.split())
            reduction = 100 * (1 - (summary_words / word_count))

            c1, c2, c3 = st.columns(3)
            c1.metric("Original Words", word_count)
            c2.metric("Summary Words", summary_words)
            c3.metric("Reduction", f"{reduction:.1f}%", delta=f"-{reduction:.1f}%")

            st.divider()

            # 2. Side-by-Side Summary Comparison
            st.markdown("Content Comparison")
            comp_col1, comp_col2 = st.columns(2)
            
            with comp_col1:
                st.markdown("**Standard Summary (S-BERT Baseline)**")
                st.info(summary_single) # Using info box for neutral baseline comparison
                st.caption(f"Word Count: {len(summary_single.split())} | Representation: Sentence-level context")

            with comp_col2:
                st.markdown("**Summary (Fused Model)**")
                st.success(summary_fused) # Using success box for proposed model distinction
                st.caption(f"Word Count: {len(summary_fused.split())} | Representation: Unified Word + Sentence + Thematic ")

            st.divider()
            
            # 3. Quality Metrics Comparison
            st.markdown("Quality Benchmarking")
            orig_center = np.mean(sent_emb, axis=0)
            def get_center(txt, full_sents, embs):
                idxs = [i for i, s in enumerate(full_sents) if s in txt]
                return np.mean(embs[idxs], axis=0) if idxs else np.zeros(embs.shape[1])

            sim_s = similarity_score(orig_center, get_center(summary_single, sentences, sent_emb))
            sim_f = similarity_score(orig_center, get_center(summary_fused, sentences, sent_emb))
            r_s = get_rouge_scores(raw_text, summary_single)
            r_f = get_rouge_scores(raw_text, summary_fused)

            q1, q2, q3 = st.columns(3)
            q1.metric("Semantic Accuracy", f"{sim_f:.4f}", f"{sim_f-sim_s:+.4f}")
            q2.metric("ROUGE-L Score", f"{r_f['rougeL'].fmeasure:.4f}", f"{r_f['rougeL'].fmeasure-r_s['rougeL'].fmeasure:+.4f}")
            q3.metric("Fusion Gain", f"{((r_f['rougeL'].fmeasure - r_s['rougeL'].fmeasure) / r_s['rougeL'].fmeasure)*100:.1f}%")

            # 4. Benchmarking Graph
            chart_data = pd.DataFrame({
                "Metric": ["Accuracy", "ROUGE-1", "ROUGE-L"],
                "Baseline (Single)": [sim_s, r_s['rouge1'].fmeasure, r_s['rougeL'].fmeasure],
                "Proposed (Fused)": [sim_f, r_f['rouge1'].fmeasure, r_f['rougeL'].fmeasure]
            })
            fig = px.bar(chart_data, x="Metric", y=["Baseline (Single)", "Proposed (Fused)"], 
                         barmode="group", color_discrete_sequence=['#A0AEC0', '#3182CE'], template="plotly_white")

            st.plotly_chart(fig, width='stretch')







